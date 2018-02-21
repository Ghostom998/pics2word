#!/usr/bin/env python3
import os, sys
from docx import Document
from docx.shared import Inches
import datetime, imghdr, struct

class PhotoCopy:

    def __init__(self):
        # set default values on instantiation, until changed with CL args
        self.SetPath()
        self.SetTitle() 
        self.SetPicWidth()
        self.SetPicHeight()
        self.GetPicsInPath()

    def SetPath(self, Path=os.getcwd()):
        # Default path is the current working directory on the command line
        self.path = Path

    def SetTitle(self, title="PhotoDoc", date='y'):
        # if date begions with 'y', append title with date
        if date[0] == 'y':
            Today = self.GetDate()
            self.title = title + "_" + str(Today)
        else:
            self.title = title

    def SetPicWidth(self,Width=4):
        # TODO set a default!
        self.width = Width

    def SetPicHeight(self,Height=4):
        # TODO set a default!
        self.height = Height

    def SetTableWidth(self, Columns=2):
        # TODO set a default!
        self.tablecolumns = Columns

    def GetDate(self):
        return datetime.date.today().strftime("%d%b%Y") # i.e. 15Feb2018

    def help(self):
        message = '''Usage: python PhotoCopy.py [-command] [value]
Options:
\t-h\t- Pass "help" to print this help message to the terminal. \n\t\t  Pass the name of a command below without the '-' for more informatio about that command. <UNDER CONSTRUCTION>
\t-P\t- Pass an alternative path to be used. i.e. \"C:\\\\Pictures\\\". Defaults to current directory.
\t-f\t- format pictures. pass either "normal" or "table". Defaults to normal. <UNDER CONSTRUCTION>
\t-T\t- Override the default title. Defaults to PhotoDoc_<current date> (See Td, below).
\t-Td\t- Choose to append the title with the current date. Options: \"y\" or \"n\". Defaults to \"y\".
\t-pw\t- Set the width of imported pictures in inches. Defaults to ........
\t-ph\t- Set the hieght of imported pictures in inches. Defaults to ........
\t-tw\t- Set the number of columns used in table format. Note: table format must be enabled! Defaults to 2.

Commands may be passed as command-value pairs in any order.
All commands are optional and the defaults will be used if no commands are given.

Example: python PhotoCopy.py -P \"C:\\\\Pictures\\\" -T \"Report\" -Td \"n\" -f \"table\"\n'''
        print(message)
    
    def Format(self, format="normal"):
        if format == "table":
            self.format = "table"
        if format == "normal":
            self.format = "normal"
        else:
            raise ValueError("Please enter a valid format for '-f' ")

    # Import files as list
    def WriteDoc(self):
        document = Document()
        p = document.add_paragraph()
        Path = self.path
        PicList = sorted(self.pics)

        for Pic in PicList:
            FullImageandPath = os.path.join(Path,Pic)
            r = p.add_run()
            # Check if portrait or landscape and call add_picture with the appropriate arguements
            # in order to set an appropriate pic size and preserve the aspect ratio
            isPortrait = self.IsPortrait(FullImageandPath)
            if isPortrait:
                r.add_picture(FullImageandPath,height=Inches(self.height))
            else:
                r.add_picture(FullImageandPath,width=Inches(self.width))
            # TODO "chop" pic name to remove extension
            p.add_run("\n"+Pic.split('.')[0]+"\n")

        document.save(self.title + '.docx')

    def WriteTable(self):
        document = Document()
        tbl = document.add_table(rows=0,cols=self.tablecolumns)
        PicList = sorted(self.pics)
        while PicList:
            row_cells = tbl.add_row().cells
            for idx, cell in enumerate(self.tablecolumns):
                if idx%2 != 0:
                    # if odd row, add picture
                    p = row_cells[idx].paragraphs[0]
                    FullImageandPath = os.path.join(self.path,PicList[0])
                    isPortrait = self.IsPortrait(FullImageandPath)
                    r = p.add_run()
                    if isPortrait:
                        r.add_picture(FullImageandPath,height=Inches(self.height))
                    else:
                        r.add_picture(FullImageandPath,width=Inches(self.width))
                    PicList = PicList[1:]
                else:
                    # else if even row, i.e. second, etc. add description


    def GetPicsInPath(self):
        self.pics = []
        ValidExtList = [".jpg",".jpeg",".png",".bmp",".gif",".JPG",".JPEG",".PNG",".BMP",".GIF"]
        for file in os.listdir(self.path):
            for ValidExt in ValidExtList:
                if file.endswith(ValidExt):
                    self.pics.append(file)
        return self.pics

    def IsPortrait(self, fname):
        """Determine the image type of fhandle and return its size."""
        with open(fname, 'rb') as fhandle:
            head = fhandle.read(24)
            if len(head) != 24:
                return
            if imghdr.what(fname) == 'png':
                check = struct.unpack('>i', head[4:8])[0]
                if check != 0x0d0a1a0a:
                    return
                width, height = struct.unpack('>ii', head[16:24])
            elif imghdr.what(fname) == 'gif':
                width, height = struct.unpack('<HH', head[6:10])
            elif imghdr.what(fname) == 'jpeg':
                try:
                    fhandle.seek(0) # Read 0xff next
                    size = 2
                    ftype = 0
                    while not 0xc0 <= ftype <= 0xcf:
                        fhandle.seek(size, 1)
                        byte = fhandle.read(1)
                        while ord(byte) == 0xff:
                            byte = fhandle.read(1)
                        ftype = ord(byte)
                        size = struct.unpack('>H', fhandle.read(2))[0] - 2
                    # We are at a SOFn block
                    fhandle.seek(1, 1)  # Skip `precision' byte.
                    height, width = struct.unpack('>HH', fhandle.read(4))
                except Exception: #IGNORE:W0703
                    return
            else:
                return
            if width/height > 1:
                return False
            else:
                return True  

    def isNumbered(self,list):
        count = 0
        for value in list:
            string = value.split('.')[0]
            if string[-1].isdigit() or string[0].isdigit:
                count += 1
        # if all names start or end with numbers, 
        # then we can assume they have been numbered
        if count == len(list):
            sorting_tuple = [()]
            for value in list:
                string = value.split('.')[0]
                string[len(string.rstrip('0123456789')):]
            return True
        else:
            return False

'''
#!/usr/bin/env python3
import os, sys
from .PhotoCopy import *
'''
# Method to parse command line arguements into command-value pairs
def getopts(argv):
    opts = {}  # Empty dictionary to store key-value pairs.
    while argv:  # While there are arguments left to parse...
        if argv[0][0] == '-':  # Found a "-name value" pair.
            opts[argv[0]] = argv[1]  # Add key and value to the dictionary.
        argv = argv[1:]  # Reduce the argument list by copying it starting from index 1.
    return opts

def main():
    Doc = PhotoCopy()
    myargs = getopts(sys.argv)
    if '-h' in myargs:
        Doc.help()
    if '-P' in myargs:
        # Override the default path
        Doc.SetPath(myargs['-P'])
    if '-f' in myargs:
        # Set as table or default
        Doc.Format(myargs['-f'])
    if '-T' in myargs:
        # Set a title to override the default
        Doc.SetTitle(title=myargs['-T'],date=myargs['-Td'])
    if '-pw' in myargs:
        # Override the default picture width
        Doc.SetPicWidth(myargs['-pw'])
    if '-ph' in myargs:
        # Override the default picture height
        Doc.SetPicHeight(myargs['-ph']) 
    if '-tw' in myargs:
        if myargs['-f'] is not "table":
            raise ValueError("Must enable table format to format table width!")
        else:
            Doc.SetTableWidth(myargs['-tw'])
    
    # after all optional parameters have been changed and not asked for help, then write document.
    if '-h' not in myargs:
        Doc.WriteDoc()

if __name__ == '__main__':
    main()
